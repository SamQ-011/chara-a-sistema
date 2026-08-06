import os
from copy import deepcopy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import _Row
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from num2words import num2words

def numero_a_letras(monto: float):
    entero = int(monto)
    decimal = round((monto - entero) * 100)
    texto = num2words(entero, lang="es").title()
    if texto.startswith("Mil "):
        texto = "Un " + texto
    elif texto == "Mil":
        texto = "Un Mil"
    return f"Bs.{monto:,.2f} ({texto} {decimal:02d}/100)"

# ==========================================================
# REEMPLAZAR TEXTO (VERSIÓN BLINDADA CONTRA XML ROTO)
# ==========================================================
def reemplazar_texto(doc, reemplazos):
    def replace_in_paragraph(p):
        if not p.runs:
            return
        texto_completo = "".join([run.text for run in p.runs])
        modificado = False
        
        for buscar, reemplazo in reemplazos.items():
            if buscar in texto_completo:
                texto_completo = texto_completo.replace(buscar, str(reemplazo))
                modificado = True
                
        if modificado:
            p.runs[0].text = texto_completo
            for run in p.runs[1:]:
                run.text = ""

    for p in doc.paragraphs:
        replace_in_paragraph(p)
        
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    replace_in_paragraph(p)


# ==========================================================
# INYECCIÓN DINÁMICA DE NODOS (VIÑETAS Y LISTAS)
# ==========================================================
def inyectar_nodos_dinamicos(doc, marcador, lista_textos):
    if not lista_textos:
        def eliminar_nodo_marcador(elementos):
            for p in elementos:
                if marcador in p.text:
                    padre = p._p.getparent()
                    padre.remove(p._p)
                    return True
            return False

        if not eliminar_nodo_marcador(doc.paragraphs):
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if eliminar_nodo_marcador(cell.paragraphs):
                            break
        return

    # Forzar a Word a que NO aplaste los párrafos del mismo estilo
    def quitar_contextual_spacing(parrafo):
        pPr = parrafo._p.get_or_add_pPr()
        existente = pPr.find(qn('w:contextualSpacing'))
        if existente is not None:
            existente.set(qn('w:val'), '0')
            return
        nuevo = OxmlElement('w:contextualSpacing')
        nuevo.set(qn('w:val'), '0')
        rPr = pPr.find(qn('w:rPr'))
        if rPr is not None:
            rPr.addprevious(nuevo)
        else:
            pPr.append(nuevo)

    def buscar_y_reemplazar_nodos(elementos):
        for p in elementos:
            if marcador in p.text:
                padre = p._p.getparent()
                indice = padre.index(p._p)
                
                # Extraer formato manual (tipografía) del marcador original
                fuente_nombre = None
                fuente_tamano = None
                if p.runs:
                    fuente_nombre = p.runs[0].font.name
                    fuente_tamano = p.runs[0].font.size

                for item in lista_textos:
                    if isinstance(item, dict):
                        # ==========================================
                        # MODO "PUNTO EXTRA" (TÍTULO Y DESCRIPCIÓN)
                        # ==========================================
                        
                        # 1. PÁRRAFO DEL TÍTULO (Mantiene la numeración)
                        p_tit = deepcopy(p._p)
                        parrafo_tit = Paragraph(p_tit, p._parent)
                        parrafo_tit.clear()
                        
                        run_tit = parrafo_tit.add_run(str(item.get('titulo', '')).strip())
                        run_tit.bold = True
                        if fuente_nombre: run_tit.font.name = fuente_nombre
                        if fuente_tamano: run_tit.font.size = fuente_tamano
                        
                        # Espaciado explícito: título pegado a su descripción
                        quitar_contextual_spacing(parrafo_tit)
                        parrafo_tit.paragraph_format.space_before = Pt(6)
                        parrafo_tit.paragraph_format.space_after = Pt(6)
                        
                        padre.insert(indice, parrafo_tit._p)
                        indice += 1
                        
                        # 2. PÁRRAFO DE LA DESCRIPCIÓN (Sin numeración, misma sangría)
                        p_desc = deepcopy(p._p)
                        parrafo_desc = Paragraph(p_desc, p._parent)
                        parrafo_desc.clear()
                        
                        # TRUCO: Eliminar la etiqueta XML de numeración (numPr) pero dejar la sangría (ind)
                        pPr = parrafo_desc._p.get_or_add_pPr()
                        for child in list(pPr):
                            if child.tag.endswith('numPr'):
                                pPr.remove(child)
                        quitar_contextual_spacing(parrafo_desc)
                                
                        run_desc = parrafo_desc.add_run(str(item.get('descripcion', '')).strip())
                        run_desc.bold = False
                        if fuente_nombre: run_desc.font.name = fuente_nombre
                        if fuente_tamano: run_desc.font.size = fuente_tamano
                        
                        # Espaciado explícito: aire antes del siguiente punto
                        parrafo_desc.paragraph_format.space_before = Pt(0)
                        parrafo_desc.paragraph_format.space_after = Pt(12)
                        
                        padre.insert(indice, parrafo_desc._p)
                        indice += 1
                        
                    else:
                        # ==========================================
                        # MODO "VIÑETA NORMAL" (Otras Condiciones)
                        # ==========================================
                        nuevo_p = deepcopy(p._p)
                        nuevo_parrafo = Paragraph(nuevo_p, p._parent)
                        nuevo_parrafo.clear()
                        
                        run_v = nuevo_parrafo.add_run(str(item).strip())
                        if fuente_nombre: run_v.font.name = fuente_nombre
                        if fuente_tamano: run_v.font.size = fuente_tamano
                        
                        # ESPACIADO: Pegamos las viñetas quitando el espacio posterior
                        quitar_contextual_spacing(nuevo_parrafo)
                        nuevo_parrafo.paragraph_format.space_before = Pt(0)
                        nuevo_parrafo.paragraph_format.space_after = Pt(0)
                        
                        padre.insert(indice, nuevo_parrafo._p)
                        indice += 1
                        
                # Destruimos el marcador original
                padre.remove(p._p)
                return True
        return False

    # Ejecutar la búsqueda en párrafos normales
    encontrado = buscar_y_reemplazar_nodos(doc.paragraphs)
    
    # Si no estaba en el cuerpo, buscar dentro de las tablas
    if not encontrado:
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    if buscar_y_reemplazar_nodos(celda.paragraphs):
                        return

# ==========================================================
# LLENAR TABLAS DINÁMICAS (INTELIGENTE Y COMPACTO)
# ==========================================================
def llenar_tablas_scp(doc, items, gastos):
    # Remueve el margen extra inferior de los párrafos en las celdas
    def set_celda(fila, index, valor):
        if index < len(fila.cells):
            celda = fila.cells[index]
            celda.text = str(valor)
            for p in celda.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)

    # Convierte flotantes como 3.0 a enteros limpios (3)
    def format_cantidad(val):
        try:
            f_val = float(val)
            if f_val.is_integer():
                return str(int(f_val))
            return str(f_val)
        except (ValueError, TypeError):
            return str(val)

    # 1. Buscar Tabla de Gastos
    tabla_gastos = None
    for t in doc.tables:
        if len(t.rows) > 0:
            header = "".join([c.text.upper() for c in t.rows[0].cells])
            if "OBJETO DEL GASTO" in header or "PROG" in header:
                tabla_gastos = t
                break
                
    # 2. Buscar Tabla de Ítems
    tabla_items = None
    for t in doc.tables:
        if len(t.rows) > 0:
            header = "".join([c.text.upper() for c in t.rows[0].cells])
            if "CANTIDAD" in header and "UNIDAD" in header:
                tabla_items = t
                break

    # Rellenar Gastos
    if gastos and tabla_gastos:
        if len(tabla_gastos.rows) > 1:
            fila_datos = tabla_gastos.rows[1]
            tiene_monto = len(fila_datos.cells) >= 8
            offset = 1 if tiene_monto else 0
            
            monto_val = f"{float(gastos[0].get('monto', 0)):,.2f}"
            if tiene_monto:
                set_celda(fila_datos, 0, monto_val)
                
            set_celda(fila_datos, offset, gastos[0].get("partida", ""))
            set_celda(fila_datos, offset + 1, gastos[0].get("prog", ""))
            set_celda(fila_datos, offset + 2, gastos[0].get("proy", ""))
            set_celda(fila_datos, offset + 3, gastos[0].get("act", ""))
            set_celda(fila_datos, offset + 4, gastos[0].get("ff", ""))
            set_celda(fila_datos, offset + 5, gastos[0].get("of", ""))
            set_celda(fila_datos, offset + 6, gastos[0].get("descripcion", ""))
            
            ultima_fila_gastos = fila_datos
            for gasto in gastos[1:]:
                nueva_fila_tr = deepcopy(fila_datos._tr)
                ultima_fila_gastos._tr.addnext(nueva_fila_tr)
                nueva_fila = _Row(nueva_fila_tr, tabla_gastos)
                
                monto_val = f"{float(gasto.get('monto', 0)):,.2f}"
                if tiene_monto:
                    set_celda(nueva_fila, 0, monto_val)
                    
                set_celda(nueva_fila, offset, gasto.get("partida", ""))
                set_celda(nueva_fila, offset + 1, gasto.get("prog", ""))
                set_celda(nueva_fila, offset + 2, gasto.get("proy", ""))
                set_celda(nueva_fila, offset + 3, gasto.get("act", ""))
                set_celda(nueva_fila, offset + 4, gasto.get("ff", ""))
                set_celda(nueva_fila, offset + 5, gasto.get("of", ""))
                set_celda(nueva_fila, offset + 6, gasto.get("descripcion", ""))
                ultima_fila_gastos = nueva_fila

    # Rellenar Ítems
    if items and tabla_items:
        if len(tabla_items.rows) >= 2:
            fila_datos = tabla_items.rows[-2]
            fila_footer = tabla_items.rows[-1]
            
            # --- CORRECCIÓN: Concatenar Objeto Corto + Desc. Larga ---
            obj_corto = str(items[0].get("objeto", "")).upper()
            desc_larga = str(items[0].get("descripcion", ""))
            texto_item = f"{obj_corto}\n{desc_larga}" if desc_larga.strip() else obj_corto
            
            set_celda(fila_datos, 0, items[0].get("nro", ""))
            set_celda(fila_datos, 1, format_cantidad(items[0].get("cant", "")))
            set_celda(fila_datos, 2, items[0].get("tipuni", ""))
            set_celda(fila_datos, 3, texto_item) # <--- Aquí se inyecta la combinación
            set_celda(fila_datos, 4, f"{float(items[0].get('precio_unitario', 0)):,.2f}")
            set_celda(fila_datos, 5, f"{float(items[0].get('total_item', 0)):,.2f}")
            
            for item in items[1:]:
                nueva_fila_tr = deepcopy(fila_datos._tr)
                fila_footer._tr.addprevious(nueva_fila_tr)
                nueva_fila = _Row(nueva_fila_tr, tabla_items)
                
                # --- CORRECCIÓN REPETIDA PARA EL BUCLE ---
                obj_corto_bucle = str(item.get("objeto", "")).upper()
                desc_larga_bucle = str(item.get("descripcion", ""))
                texto_item_bucle = f"{obj_corto_bucle}\n{desc_larga_bucle}" if desc_larga_bucle.strip() else obj_corto_bucle
                
                set_celda(nueva_fila, 0, item.get("nro", ""))
                set_celda(nueva_fila, 1, format_cantidad(item.get("cant", "")))
                set_celda(nueva_fila, 2, item.get("tipuni", ""))
                set_celda(nueva_fila, 3, texto_item_bucle)
                set_celda(nueva_fila, 4, f"{float(item.get('precio_unitario', 0)):,.2f}")
                set_celda(nueva_fila, 5, f"{float(item.get('total_item', 0)):,.2f}")

def llenar_tabla_certificacion(doc, gastos):
    if not gastos:
        return
        
    tabla_obj = None
    fila_datos_idx = -1
    
    for t in doc.tables:
        for idx, fila in enumerate(t.rows):
            texto_fila = "".join([c.text.upper() for c in fila.cells])
            if "APERTURA PROGRAMATICA" in texto_fila and "FUENTE Y ORG" in texto_fila:
                tabla_obj = t
                fila_datos_idx = idx + 1
                break
        if tabla_obj:
            break
            
    if tabla_obj and fila_datos_idx < len(tabla_obj.rows):
        fila_datos = tabla_obj.rows[fila_datos_idx] 
        
        def set_celda(fila, index, valor):
            if index < len(fila.cells):
                celda = fila.cells[index]
                celda.text = str(valor)
                for p in celda.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)

        ultima_fila = fila_datos
        
        for i, gasto in enumerate(gastos):
            if i == 0:
                fila_actual = fila_datos
            else:
                nueva_fila_tr = deepcopy(fila_datos._tr)
                ultima_fila._tr.addnext(nueva_fila_tr)
                fila_actual = _Row(nueva_fila_tr, tabla_obj)
                ultima_fila = fila_actual
            
            ap_nuevo = f"{gasto.get('prog', '')} {gasto.get('proy', '')} {gasto.get('act', '')}"
            ff = str(gasto.get('ff', ''))
            of = str(gasto.get('of', ''))
            fo_nuevo = f"{ff}/{of}" if (ff and of) else "S/N"
            
            # NOTA: la plantilla tiene 2 celdas fusionadas (gridSpan) para 5 columnas visibles = 7 posiciones reales en fila.cells:
            #   - "FUENTE Y ORG. FINANCIADOR" ocupa los índices 1 y 2 (misma celda)
            #   - "DESCRIPCIÓN" ocupa los índices 4 y 5 (misma celda)
            # Por eso cada tramo se escribe UNA sola vez, en el último índice de cada par fusionado.
            set_celda(fila_actual, 0, ap_nuevo)                                # Apertura Programática
            set_celda(fila_actual, 2, fo_nuevo)                                 # Fuente y Org. Financiador -> FF/OF (fusionada 1-2)
            set_celda(fila_actual, 3, str(gasto.get("partida", "")))            # Objeto de Gasto -> Partida
            set_celda(fila_actual, 5, str(gasto.get("descripcion", "")))       # Descripción (fusionada 4-5)
            set_celda(fila_actual, 6, f"{float(gasto.get('monto', 0)):,.2f}")  # Presupuesto Solicitado

def llenar_tabla_solicitud_inicio(doc, items):
    if not items:
        return

    def format_cantidad(val):
        try:
            f_val = float(val)
            return str(int(f_val)) if f_val.is_integer() else str(f_val)
        except: return str(val)

    def set_celda(fila, index, valor):
        if index < len(fila.cells):
            celda = fila.cells[index]
            celda.text = str(valor)
            for p in celda.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)

    # Nuevo motor multilinea que imita el "new Break()" de C#
    def set_celda_multilinea(fila, index, obj_corto, desc_larga):
        if index < len(fila.cells):
            celda = fila.cells[index]
            celda.text = "" 
            p = celda.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            
            run_base = p.add_run(str(obj_corto).upper())
            
            desc_str = str(desc_larga).strip()
            if desc_str:
                run_base.add_break() 
                lineas = desc_str.split('\n')
                for i, linea in enumerate(lineas):
                    p.add_run(linea)
                    if i < len(lineas) - 1:
                        p.add_run().add_break()

    tabla_items = None
    for t in doc.tables:
        if len(t.rows) > 0:
            header = "".join([c.text.upper() for c in t.rows[0].cells])
            # CORRECCIÓN 1: Detección más flexible para los nuevos títulos de la plantilla
            if "CANTIDAD" in header and "UNIDAD" in header:
                tabla_items = t
                break

    if tabla_items and len(tabla_items.rows) >= 2:
        fila_datos = tabla_items.rows[-2] 
        fila_footer = tabla_items.rows[-1] 
        
        for i, item in enumerate(items):
            if i == 0:
                fila_actual = fila_datos
            else:
                nueva_fila_tr = deepcopy(fila_datos._tr)
                fila_footer._tr.addprevious(nueva_fila_tr)
                fila_actual = _Row(nueva_fila_tr, tabla_items)
            
            obj_c = item.get("objeto", "")
            desc_l = item.get("descripcion", "")
            
            set_celda(fila_actual, 0, item.get("nro", ""))
            set_celda_multilinea(fila_actual, 1, obj_c, desc_l) 
            
            # CORRECCIÓN 2: Inversión de columnas (Índice 2 = Unidad, Índice 3 = Cantidad)
            set_celda(fila_actual, 2, item.get("tipuni", ""))
            set_celda(fila_actual, 3, format_cantidad(item.get("cant", "")))
            
            set_celda(fila_actual, 4, f"{float(item.get('precio_unitario', 0)):,.2f}")
            set_celda(fila_actual, 5, f"{float(item.get('total_item', 0)):,.2f}")

def llenar_tabla_informe_cotizacion(doc, items):
    if not items:
        return

    def set_celda(fila, index, valor):
        if index < len(fila.cells):
            celda = fila.cells[index]
            celda.text = str(valor)
            for p in celda.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)

    # 1. Buscar la tabla que contenga 'PROVEEDOR' en su encabezado
    tabla_cot = None
    for t in doc.tables:
        if len(t.rows) > 0:
            header = "".join([c.text.upper() for c in t.rows[0].cells])
            if "PROVEEDOR" in header and "PRECIO" in header:
                tabla_cot = t
                break

    # 2. Si encuentra la tabla, llenar iterando sobre 'items' (cotizaciones)
    if tabla_cot and len(tabla_cot.rows) >= 2:
        fila_datos = tabla_cot.rows[1] # Fila base que contiene 'n', 'pr', 'dec', 'pt'
        ultima_fila = fila_datos

        for i, item in enumerate(items):
            if i == 0:
                fila_actual = fila_datos
            else:
                nueva_fila_tr = deepcopy(fila_datos._tr)
                ultima_fila._tr.addnext(nueva_fila_tr)
                fila_actual = _Row(nueva_fila_tr, tabla_cot)
                ultima_fila = fila_actual
            
            set_celda(fila_actual, 0, item.get("n", str(i + 1)))
            set_celda(fila_actual, 1, str(item.get("pr", "")))
            set_celda(fila_actual, 2, str(item.get("dec", "")))
            
            try:
                precio_fmt = f"{float(item.get('pt', 0)):,.2f}"
            except Exception:
                precio_fmt = str(item.get("pt", "0.00"))
            
            set_celda(fila_actual, 3, precio_fmt)

def llenar_tabla_acta_entrega(doc, items):
    if not items:
        return

    def set_celda(fila, index, valor):
        if index < len(fila.cells):
            celda = fila.cells[index]
            celda.text = str(valor)
            for p in celda.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)

    # 1. Buscar la tabla del acta (buscamos 'OBSERVACIONES' en la cabecera)
    tabla_acta = None
    for t in doc.tables:
        if len(t.rows) > 0:
            header = "".join([c.text.upper() for c in t.rows[0].cells])
            if "OBSERVACIONES" in header or "CANT" in header:
                tabla_acta = t
                break

    # 2. Rellenar clonando la fila 1 (donde están nro, desc, cant, unm)
    if tabla_acta and len(tabla_acta.rows) >= 2:
        fila_datos = tabla_acta.rows[1]
        ultima_fila = fila_datos
        
        for i, item in enumerate(items):
            if i == 0:
                fila_actual = fila_datos
            else:
                nueva_fila_tr = deepcopy(fila_datos._tr)
                ultima_fila._tr.addnext(nueva_fila_tr)
                fila_actual = _Row(nueva_fila_tr, tabla_acta)
                ultima_fila = fila_actual
            
            # Mapeo según la tabla enviada desde generador_documentos.py
            set_celda(fila_actual, 0, item.get("nro", ""))
            set_celda(fila_actual, 1, item.get("desc", ""))
            set_celda(fila_actual, 2, item.get("cant", ""))
            set_celda(fila_actual, 3, item.get("unm", ""))
            set_celda(fila_actual, 4, "")

def eliminar_fila_retencion(doc):
    """Busca dinámicamente la fila de retenciones y la elimina desde el XML"""
    for table in doc.tables:
        for row in table.rows:
            # Unimos todo el texto de la fila y lo pasamos a minúsculas para buscar
            texto_fila = "".join(cell.text for cell in row.cells).lower()
            
            # Si encuentra la palabra clave de esa fila en el Word, la elimina
            if "retención" in texto_fila or "retencion" in texto_fila or "retenciones" in texto_fila:
                tbl = table._tbl
                tr = row._tr
                tbl.remove(tr)
                return

def insertar_documentos(doc, documentos):
    if not documentos:
        return
    for paragraph in doc.paragraphs:
        if "{DOCUMENTOS}" in paragraph.text:
            padre = paragraph._p.getparent()
            indice = padre.index(paragraph._p)
            for texto in documentos:
                nuevo_xml = deepcopy(paragraph._p)
                nuevo = Paragraph(nuevo_xml, paragraph._parent)
                if nuevo.runs:
                    nuevo.runs[0].text = texto
                    for run in nuevo.runs[1:]:
                        run.text = ""
                else:
                    nuevo.add_run(texto)
                padre.insert(indice, nuevo._p)
                indice += 1
            padre.remove(paragraph._p)
            return

def llenar_tabla_informe_conformidad(doc, items):
    if not items:
        return

    def set_celda(fila, index, valor):
        if index < len(fila.cells):
            celda = fila.cells[index]
            celda.text = str(valor)
            for p in celda.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)

    tabla_info = None
    for t in doc.tables:
        if len(t.rows) > 0:
            header = "".join([c.text.upper() for c in t.rows[0].cells])
            if "PRECIO TOTAL" in header and "UNIDAD DE" in header:
                tabla_info = t
                break

    if tabla_info and len(tabla_info.rows) >= 2:
        fila_datos = tabla_info.rows[1]
        ultima_fila = fila_datos
        
        for i, item in enumerate(items):
            if i == 0:
                fila_actual = fila_datos
            else:
                nueva_fila_tr = deepcopy(fila_datos._tr)
                ultima_fila._tr.addnext(nueva_fila_tr)
                fila_actual = _Row(nueva_fila_tr, tabla_info)
                ultima_fila = fila_actual
            
            obj = str(item.get("objeto", "")).strip()
            desc = str(item.get("descripcion", "")).strip()
            texto_unido = f"{obj}\n{desc}" if desc else obj
            
            set_celda(fila_actual, 0, item.get("nro", ""))
            set_celda(fila_actual, 1, texto_unido)
            set_celda(fila_actual, 2, str(item.get("cant", "")))
            set_celda(fila_actual, 3, item.get("tipuni", ""))
            set_celda(fila_actual, 4, f"{float(item.get('total_item', 0)):,.2f}")

def generar_documento_word(
    ruta_plantilla, ruta_salida, fecha="", cod="", nombre="", fechaInfo="", 
    retenc=False, porcentaje=0, objContr="", monto=0, plazoEntrega=0, 
    seleccionados=None, items=None, gastos=None, variables_extra=None, 
    lotes_actas=None, puntos_extra=None, condiciones=None
):
    if not os.path.exists(ruta_plantilla):
        raise FileNotFoundError(f"Plantilla no encontrada: {ruta_plantilla}")

    os.makedirs(os.path.dirname(ruta_salida), exist_ok=True)
    doc = Document(ruta_plantilla)

    reemplazos = {
        "{FECHA}": fecha,
        "{COD}": cod,
        "{NOMBRE}": nombre,
        "{FECHAINF}": fechaInfo,
        "{OBJCONTR}": objContr.upper() if objContr else "",
        "{MONTO}": numero_a_letras(monto),
        "{RETENC}": numero_a_letras(porcentaje),
        "{PE}": f"{plazoEntrega} días Hábiles a partir de la Recepción de la Orden de Compra"
    }

    if variables_extra:
        reemplazos.update(variables_extra)

    if "solicitudCP" in ruta_plantilla:
        llenar_tablas_scp(doc, items, gastos)
    elif "certificacionPresupuestaria" in ruta_plantilla:
        llenar_tabla_certificacion(doc, gastos)
    elif "solicitudIniProcContr" in ruta_plantilla or "espTec" in ruta_plantilla:
        llenar_tabla_solicitud_inicio(doc, items)
    elif "infoCotizacion" in ruta_plantilla:
        llenar_tabla_informe_cotizacion(doc, items)
    elif "actaEntrega" in ruta_plantilla and lotes_actas:
        doc_master = None
        for lote_items in lotes_actas:
            # 1. Instanciamos una plantilla fresca por cada acta
            doc_temp = Document(ruta_plantilla)
            reemplazar_texto(doc_temp, reemplazos)
            llenar_tabla_acta_entrega(doc_temp, lote_items)
            
            # 2. Si es la primera hoja, se vuelve nuestro Master
            if doc_master is None:
                doc_master = doc_temp
            else:
                # 3. Si es la 2da, 3ra hoja, insertamos salto de página y clonamos el XML puro
                doc_master.add_page_break()
                for element in doc_temp.element.body:
                    if element.tag.endswith('sectPr'): # Ignoramos configuración de página oculta para no corromper
                        continue
                    doc_master.element.body.append(deepcopy(element))
        
        # Guardamos el archivo unificado gigante y terminamos
        if doc_master:
            doc_master.save(ruta_salida)
            return
            
    elif "infoConformidad" in ruta_plantilla:
        llenar_tabla_informe_conformidad(doc, items)

    if puntos_extra is not None:
        inyectar_nodos_dinamicos(doc, "{PUNTOS_EXTRA}", puntos_extra)
    
    if condiciones is not None:
        inyectar_nodos_dinamicos(doc, "{VIÑETAS_CONDICIONES}", condiciones)
    
    # Ejecución final limpia (sin duplicados)
    if "actaEntrega" not in ruta_plantilla:
        reemplazar_texto(doc, reemplazos)
        
        # Si NO hay retención y es la notificación, borramos la fila
        if not retenc and "notificAdjudic" in ruta_plantilla:
            eliminar_fila_retencion(doc)
            
        insertar_documentos(doc, seleccionados)
        doc.save(ruta_salida)